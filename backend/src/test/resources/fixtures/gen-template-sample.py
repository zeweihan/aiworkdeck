#!/usr/bin/env python3
"""生成 DocxProfileReaderTest 用的模板样本 template-sample.docx。

刻意造出与 HOUSE 不同的一套版式，让读端每个字段都有区分度：
- 正文 Normal：楷体_GB2312 / Arial 12pt，两端对齐，段后 18pt，行距最小值 16pt，无首行缩进
- Heading 1-3：自动编号（numbering.xml）一、（一）1.；H1 黑体 16pt 粗居中，H2/H3 楷体 12pt 粗
- 一张 3x3 表：单元格级边框（tcBorders 0.5pt）、gridCol 2019/3500/3507、表头粗+底纹 D9D9D9+重复表头、
  10pt 字、数字居右、垂直居中
- 目录域 TOC \\o "1-2" \\h \\z \\u
- 页脚「第 PAGE 页」居中
- A4，上下 2.54cm，左右 3.17cm

用法：python3 gen-template-sample.py  （在本目录运行，覆盖 template-sample.docx）
依赖：python-docx >= 1.1
"""
import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "template-sample.docx")


def set_fonts(rpr_owner_font, east_asia, western):
    font = rpr_owner_font
    font.name = western
    rpr = font.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), western)
    rfonts.set(qn("w:hAnsi"), western)
    rfonts.set(qn("w:cs"), western)
    rfonts.set(qn("w:eastAsia"), east_asia)
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]


def clear_heading_color(style):
    rpr = style.element.rPr
    if rpr is None:
        return
    color = rpr.find(qn("w:color"))
    if color is not None:
        rpr.remove(color)


def add_numbering(doc):
    """建一个 abstractNum（三级 一、/（一）/1.）并返回 numId。"""
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(a.get(qn("w:abstractNumId"))) for a in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abs_id = (max(abstract_ids) + 1) if abstract_ids else 0
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    levels = [
        ("chineseCountingThousand", "%1、", "nothing", "Heading1"),
        ("chineseCountingThousand", "（%2）", "nothing", "Heading2"),
        ("decimal", "%3.", "space", "Heading3"),
    ]
    for ilvl, (fmt, text, suff, pstyle) in enumerate(levels):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), pstyle)
        lvl.append(ps)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), suff)
        lvl.append(suffix)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:firstLine"), "0")
        ppr.append(ind)
        lvl.append(ppr)
        abstract.append(lvl)

    # abstractNum 必须排在所有 num 之前
    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract)
    else:
        numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abs_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def style_num_pr(style, num_id, ilvl):
    ppr = style.element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    il = OxmlElement("w:ilvl")
    il.set(qn("w:val"), str(ilvl))
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    num_pr.append(il)
    num_pr.append(nid)
    ppr.append(num_pr)


def set_para_format(pf, align, before_pt, after_pt, at_least_pt, first_line_cm):
    pf.alignment = align
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    pf.line_spacing = Pt(at_least_pt)
    pf.first_line_indent = Cm(first_line_cm)


def add_field(paragraph, instr):
    """fldChar begin / instrText / separate / 占位文本 / end 五件套。"""
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    r = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r._r.append(it)
    r = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)
    paragraph.add_run("1")
    r = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def cell_borders(cell, sz_eighths):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz_eighths))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tc_pr.append(borders)


def cell_shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_valign_center(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tc_pr.append(v)


def main():
    doc = Document()

    # 页面：A4，上下 2.54cm，左右 3.17cm
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    # Normal
    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    set_fonts(normal.font, "楷体_GB2312", "Arial")
    set_para_format(normal.paragraph_format, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 18, 16, 0)

    num_id = add_numbering(doc)

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    set_fonts(h1.font, "黑体", "Arial")
    clear_heading_color(h1)
    set_para_format(h1.paragraph_format, WD_ALIGN_PARAGRAPH.CENTER, 12, 12, 16, 0)
    style_num_pr(h1, num_id, 0)

    for name, ilvl in (("Heading 2", 1), ("Heading 3", 2)):
        h = doc.styles[name]
        h.font.size = Pt(12)
        h.font.bold = True
        h.font.italic = False
        set_fonts(h.font, "楷体_GB2312", "Arial")
        clear_heading_color(h)
        set_para_format(h.paragraph_format, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 18, 16, 0)
        style_num_pr(h, num_id, ilvl)

    # 目录
    toc_title = doc.add_paragraph("目  录")
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_p = doc.add_paragraph()
    add_field(toc_p, ' TOC \\o "1-2" \\h \\z \\u ')

    # 正文骨架：三级标题各若干 + 正文段
    for i in range(3):
        doc.add_paragraph("一级标题样例", style="Heading 1")
        doc.add_paragraph("本段为正文样例，用于统计正文段落的实例格式。" * 2)
        for j in range(2):
            doc.add_paragraph("二级标题样例", style="Heading 2")
            doc.add_paragraph("本段为正文样例，用于统计正文段落的实例格式。")
            for k in range(2):
                doc.add_paragraph("三级标题样例", style="Heading 3")
                doc.add_paragraph("本段为正文样例。")

    # 表格 3x3
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [2019, 3500, 3507]
    grid = table._tbl.tblGrid
    for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
        gc.set(qn("w:w"), str(w))
    data = [
        ["序号", "项目", "金额（万元）"],
        ["1", "注册资本", "1,000.00"],
        ["2", "实缴资本", "800.00"],
    ]
    for r, row in enumerate(table.rows):
        if r == 0:
            tr_pr = row._tr.get_or_add_trPr()
            hdr = OxmlElement("w:tblHeader")
            tr_pr.append(hdr)
        for c, cell in enumerate(row.cells):
            cell.width = Pt(widths[c] / 20)
            cell_borders(cell, 4)
            cell_valign_center(cell)
            if r == 0:
                cell_shade(cell, "D9D9D9")
            p = cell.paragraphs[0]
            run = p.add_run(data[r][c])
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(2.4)
            p.paragraph_format.space_after = Pt(2.4)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            p.paragraph_format.line_spacing = Pt(12)
            p.paragraph_format.first_line_indent = Cm(0)
            if r == 0:
                run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif c == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif c == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    after = doc.add_paragraph("表后首段正文。")
    after.paragraph_format.space_before = Pt(18)

    # 页脚页码
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("第 ")
    add_field(footer_p, " PAGE ")
    footer_p.add_run(" 页")

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Read a document into the deterministic half of a semantic map.

The pipeline has nine steps. Five of them are arithmetic and live here; three
need a model to read meaning and are specified in prompts/, and one is the
render. Keeping them apart matters: everything in this file gives the same
answer every time and can be guarded, so when a figure turns out wrong the
question is only ever whether the model's three steps were right.

    1  normalise      strip page furniture, rejoin broken lines      HERE
    2  split          sentences, with fragments folded back          HERE
    3  reconcile      sentence count == events + non-events          HERE
    4  classify       is this sentence an event?                     model
    5  date           how precise is the time, and what is it?       model
    6  segment        which part is the trunk, which the detail      model
    7  verify         every quote is findable in the normalised text HERE
    8  order          exact dates sort; relative ones by anchor      HERE
    9  render         the renderers

Every rule below was put there by a real document breaking on it. The comments
say which, because a rule whose reason has been forgotten gets deleted by the
next person who finds it inconvenient.
"""
import json, os, re, sys
from datetime import date

# ---------------------------------------------------------------- 1 normalise

RULES = [
    # A page number sitting alone on its line. In one 32-page judgment 24 of the
    # 32 page markers had been swallowed INTO a word by the extractor, so the
    # marker has to go before anything tries to read the sentence around it.
    # 页码标记 -12- 那条规则**必须排除日期形式**。
    # 它原来是 -\s*\d+\s*- ，撞上聊天记录里的 2023-04-12 就把 -04- 当页码删掉，
    # 日期变成 202312 —— 而日期是时间轴的全部信息，删掉它等于把材料毁了。
    # 判据：页码的两侧不会紧贴数字（真页码是 「\n-12-\n」这种独占一行的形式），
    # 所以用 (?<!\d) 与 (?!\d) 把「前后都是数字」的情形排除在外。
    ("page_marker",
     re.compile(r"\n?\s*(?<![\d/.])-\s*\d+\s*-(?![\d/.])\s*\n?"), "\n"),
    # 页码也可能是**裸数字单独占一行**，没有任何符号包着它。沈阳那份判决书就是这样，
    # 十处里有八处没被上一条清掉，后果不是多几个字，是改坏了材料里的数字：账号
    # 2050.01.01.01.4030 被页码 4 顶成 42050.01.01.01.4030，《情况说明》被页码 5
    # 劈成《情5况说明》。逐字核验对这种错是无效的 —— 核验比的是规范化后的文本，
    # 而错误就发生在规范化里面。
    # 这一条必须排在 soft_break 之前：先把整行页码摘掉，剩下的两半才接得回去。
    # 只认 1 到 3 位、且整行只有它。判决书里不会有一行只写一个数字的正文。
    ("page_bare", re.compile(r"(?:(?<=\n)|\A)[ \t\u3000]*\d{1,3}[ \t\u3000]*(?=\n|\Z)"), ""),
    # 「3/12」这种页码同样要排除日期：2023/4/12 里的 2023/4 会被它删掉，只剩 /12。
    # 判据：页码是**独占一行**的（前后是换行或首尾），而日期总在句子里；再排除
    # 「后面还跟着 /数字」的情形（那说明它是 年/月/日 的前两段）。
    ("page_slash",
     re.compile(r"(?:(?<=\n)|\A)\s*\d{1,3}\s*/\s*\d{1,3}\s*(?=\n|\Z)"), "\n"),
    ("form_feed", re.compile(r"\x0c"), "\n"),
    # 页码夹在日期中间时，清掉它会留下两个换行，接回断行那条规则只认单个换行、接不上，
    # 日期就永久断成两句（真材料：「2021 年 6 月 \n 5 \n 22 日签收」）。
    # 但**不能全局并空行** —— 空行是段落边界，全并掉会让段落数从 26 掉到更少，
    # 叙述块的定位跟着错（试过：句数从 46 掉到 41）。
    # 所以只在「上一行以年月结尾、下一行以日开头」这一种情形下并，判据窄到只认这件事。
    ("date_across_pages",
     re.compile(r"(\d\s*月)\s*\n\s*\n\s*(\d{1,2}\s*日)"), r"\1 \2"),
    # A line break inside a sentence. Must cover CJK<->digit in BOTH directions:
    # requiring CJK on both sides left 「…数月。」2017\n年 7 月 24 日…」 split in two
    # and made that date invisible to every later step.
    # 接回句中断行。但**下一行以「日期 + 时刻」或「说话人：」开头时，那个换行是真边界**
    # ——聊天记录与录音稿一行一条、没有句末标点，原来这条规则把十五条聊天记录粘成了
    # 三句，日期与说话人全糊在一起。而这类材料正是律师最常给的（案子多半还没判，
    # 手上是合同、聊天记录、录音稿）。所以要在接回之前把这两种行首排除掉。
    ("soft_break", re.compile(
        r"(?<=[\u4e00-\u9fff\u3001\u3002\uff09\u300b\u201d0-9])\n"
        # 判据要收紧成「**完整的日期 + 时刻**」，那才是聊天记录与录音稿的行首形式。
        # 第一版只要求「下一行以数字开头或含冒号」，于是判决书里被排版断开的
        # 「被告于 2021 年 6 月 \n 22 日签收」也被当成新的一行 —— 那是同一个日期被劈成
        # 两句，而日期是时间轴的全部信息。宽判据在这里比不判更糟。
        r"(?!\s*(?:19|20)\d{2}\s*[-/.]\s*\d{1,2}\s*[-/.]\s*\d{1,2}"
        r"\s+\d{1,2}\s*[:：]\s*\d{2})"                      # 下一行是「日期 时刻」
        r"(?!\s*[^\n：:]{1,8}[：:]\s*[^\n]{0,4}$)"            # 下一行整行是「说话人：短句」
        r"(?=[\u4e00-\u9fff0-9\uff08\u300a\u201c])"), ""),
    ("spaces", re.compile(r"[ \t\u3000]+"), " "),
    ("blank_runs", re.compile(r"\n{3,}"), "\n\n"),
]


def normalise(raw):
    text, counts = raw, {}
    for name, pat, rep in RULES:
        text, n = pat.subn(rep, text)
        counts[name] = n
    return text, counts


# --------------------------------------------------------------- 1b paragraphs

# A litigation PDF has no paragraph breaks. Of 832 lines in one judgment only 70
# ended in a full stop: every other line break is the typesetter wrapping at a
# fixed width. So a paragraph is not "text between newlines" — it is a run of
# wrapped lines ending at sentence punctuation. Looking for paragraph markers in
# the raw newlines finds nothing, which is why block detection failed three
# times before this was understood.
_PARA_END = ("\u3002", "\uff01", "\uff1f")


def paragraphs(raw):
    t = re.sub(r"\n?\s*-\s*\d+\s*-\s*\n?", "\n", raw).replace("\x0c", "")
    out, cur = [], []
    for line in (l.rstrip() for l in t.split("\n")):
        if not line.strip():
            if cur:
                out.append("".join(cur)); cur = []
            continue
        ln = line.strip()
        # A short line that does not end mid-clause stands alone: the court name,
        # the case number, a party line. Without this the whole masthead is glued
        # onto the first real paragraph and the block markers inside it end up
        # thirty characters deep, past where any opener would look.
        short_standalone = (len(ln) <= 25
                            and not ln.endswith(("\uff0c", "\u3001", "\uff1b", "\uff1a")))
        if short_standalone and not cur:
            out.append(ln)
            continue
        cur.append(ln)
        if line.endswith(_PARA_END):
            out.append("".join(cur)); cur = []
    if cur:
        out.append("".join(cur))
    return [p for p in out if p.strip()]


# Block markers. The findings section is the awkward one: judgments differ on
# whether it is announced at all. One of the two real judgments has no 「本院查明」
# anywhere, and its findings simply start after the defence and run until the
# fixed closing line 「以上事实……在案佐证」. So the findings are located by their
# END, not their beginning — the ending is standard even when the heading is not.
_OPENERS = [
    # 原告那一段并不总带「诉称」。沈阳这份写的是「原告向本院提出诉讼请求：1.判令……」，
    # 通篇没有「诉称」二字，于是整个 claim 块识别不到，而它恰是诉请金额与计息起点的
    # 出处。两种写法都要认。
    ("claim",    re.compile(r"^.{0,16}?(起诉称|诉称|向本院提出诉讼请求|诉讼请求[:：])")),
    ("defence",  re.compile(r"^.{0,16}?(共同答辩称|答辩称|辩称)")),
    # Findings may be announced three ways, and the third has no 「查明」 in it at
    # all: 「当事人围绕诉讼请求依法提交了证据……根据当事人陈述和经审查确认的证据」 is
    # boilerplate that opens the findings whenever the heading is missing. Both
    # real judgments are covered only once this is included.
    ("findings", re.compile(r"^(经审理查明|本院查明|经审理|当事人围绕诉讼请求)")),
    ("reasoning", re.compile(r"^.{0,8}?本院认为")),
    ("order",    re.compile(r"^判决如下")),
]
_FINDINGS_END = re.compile(r"^以上事实|在案佐证")


def blocks(paras):
    """Map each paragraph index to a narrative block, or None.

    Returns {block: (start, end)}. A block that cannot be located is simply
    absent — better an honest gap than a guessed boundary.
    """
    marks = {}
    for i, p in enumerate(paras):
        for name, pat in _OPENERS:
            if name not in marks and pat.match(p):
                marks[name] = i
    end_f = next((i for i, p in enumerate(paras) if _FINDINGS_END.match(p)), None)
    out = {}
    keys = sorted(marks.items(), key=lambda kv: kv[1])
    for j, (name, a) in enumerate(keys):
        b = keys[j + 1][1] if j + 1 < len(keys) else len(paras)
        out[name] = (a, b)
    if "findings" not in out and end_f is not None:
        a = marks.get("defence")
        if a is not None:
            # findings run from after the defence to the closing line
            nxt = min([v for v in marks.values() if v > a] + [end_f])
            out["findings"] = (nxt if nxt < end_f else a + 1, end_f + 1)
            if "defence" in out:
                out["defence"] = (out["defence"][0], out["findings"][0])
    return out


# -------------------------------------------------------------------- 2 split

_END = re.compile(r"(?<=[\u3002\uff01\uff1f])")
_MIN = 6


def read_with_blocks(path):
    """一次读完，句子与叙述块**出自同一份规范化文本**，句号可以直接互用。

    此前两条路径的输入不同：句子来自 normalise(raw)（26 段），而调用方习惯拿
    paragraphs(raw)（33 段）去认叙述块 —— 段号根本不可比，换算出来的句号整体偏移。
    这不是换算函数的 bug，是上游两条路径没对齐；所以修在源头，只提供这一个入口。
    """
    raw = read_docx(path) if str(path).lower().endswith(".docx") else \
        open(path, encoding="utf-8").read()
    text, counts = normalise(raw)
    paras = paragraphs(text)
    sents = split_sentences(text)
    blks = blocks(paras)
    return {"path": path, "normalised": text, "paragraphs": paras,
            "sentences": sents, "blocks": blks,
            "blocks_sids": blocks_as_sids(text, blks),
            "normalise_counts": counts}


def blocks_as_sids(text, blks):
    """把叙述块的边界从段落号换成句号。

    做法：**拿真实的切句结果去反查**，不自己再算一遍每段有几句。
    自己算过两版都对不齐：第一版按原始换行切段（61 段对 33 段），第二版复用了
    paragraphs 但折叠短句的规则与 split_sentences 差一点，全文 45 句对 46 句，
    于是「被告辩称」那一块的起点偏了一句。
    折叠规则藏在 split_sentences 里，与其抄一遍（抄错就偏），不如直接**按段落的
    首句去真实句子表里找位置** —— 同一份文本、同一套切法，不可能偏。
    """
    paras = paragraphs(text)
    sents = split_sentences(text)
    # 每段的第一句在句子表里的位置：逐段推进地找，避免同样开头的段互相抢
    starts, k = [], 0
    for para in paras:
        head = [x.strip() for x in _END.split(para) if x.strip()]
        head = head[0] if head else para
        hit = None
        for j in range(k, len(sents)):
            if sents[j].startswith(head[:12]) or head.startswith(sents[j][:12]):
                hit = j
                break
        if hit is None:
            hit = k
        starts.append(hit)
        k = hit + 1
    starts.append(len(sents))
    out = {}
    for name, (a, b) in (blks or {}).items():
        a = max(0, min(a, len(paras) - 1))
        b = max(a + 1, min(b, len(paras)))
        out[name] = (starts[a], starts[b] if b < len(starts) else len(sents))
    return out


def split_sentences(text):
    out = []
    for para in [p.strip() for p in text.split("\n") if p.strip()]:
        _first_in_para = True
        for s in [x.strip() for x in _END.split(para) if x.strip()]:
            # A fragment shorter than a few characters is punctuation debris, not
            # a sentence; fold it back so the count stays honest.
            # **但折叠不许跨段落。** 「证据清单。」这类小标题自己就是一段、只有五个字，
            # 被折回上一段的末尾之后，溯源索引里出现了「…被告仍未付款。证据清单」
            # —— 摘录里混进了下一段的标题。段落是作者给的硬边界，折叠只在段内做。
            if out and len(s) < _MIN and not _first_in_para:
                out[-1] += s
            else:
                out.append(s)
            _first_in_para = False
    return out


# ---------------------------------------------------------------- 3 reconcile

def reconcile(sentences, events, residuals):
    """Every sentence must be accounted for, as an event or as a non-event.

    This is the whole recall guarantee. Without it a model that quietly skipped
    forty sentences would look exactly like one that read them and found nothing.
    """
    seen = set()
    for e in events:
        for k in e.get("sentence_ids", []):
            seen.add(k)
    for r in residuals:
        for k in r.get("sentence_ids", []):
            seen.add(k)
    missing = [i for i in range(len(sentences)) if i not in seen]
    extra = [i for i in seen if i >= len(sentences)]
    return missing, extra


# ------------------------------------------------------------------- 7 verify

def verify_quotes(sentences, events):
    """Each quote must appear verbatim in the normalised text.

    Compared against the NORMALISED text, not the raw extraction: the raw text
    still has the line breaks and page markers that normalisation removed, so a
    quote taken from a clean sentence would never be found in it.
    """
    hay = "\n".join(sentences)
    bad = []
    for e in events:
        q = (e.get("source") or {}).get("quote")
        if not q:
            continue
        if (e.get("source") or {}).get("medium") == "image":
            continue          # read from a scan; there is no text to diff against
        if re.sub(r"\s+", "", q) not in re.sub(r"\s+", "", hay):
            bad.append((e.get("id"), q[:28]))
    return bad


# -------------------------------------------------------------------- 8 order

def _d(s):
    return date(*map(int, s.replace("-", "/").split("/")))


def solve_order(events):
    """Sort by date; report contradictions instead of silently fixing them.

    An evidence list once dated the tender submission after the award notice.
    Sorting that away would have hidden a real inconsistency in the material, so
    ordering reports and lets the lawyer decide.
    """
    def key(e):
        t = e.get("time", {})
        v = t.get("date") or t.get("from")
        return _d(v) if v else date(9999, 1, 1)

    ordered = sorted(events, key=key)
    issues = []
    for i in range(len(ordered) - 1):
        a, b = ordered[i], ordered[i + 1]
        if a.get("logical_after") == b.get("id"):
            issues.append(f"{a['id']} 依材料应在 {b['id']} 之后，但日期更早")
    anchors = {e["id"] for e in events}
    for e in events:
        t = e.get("time", {})
        if t.get("anchor") and t["anchor"] not in anchors:
            issues.append(f"{e['id']} 的锚点 {t['anchor']} 不存在")
    return ordered, issues


# ---------------------------------------------------------------------- entry

class DocxUnavailable(RuntimeError):
    """读 .docx 时缺 python-docx。单独一个类型，好让调用方决定是跳过还是报错。"""


def read_docx(path):
    """读 .docx：段落按顺序取，**表格单独处理**，不与正文混排。

    表格是这条路的关键。判决书与证据目录里的表格，用通用文本抽取（pdftotext 之类）
    会把「编号 / 名称 / 证明目的 / 页码」挤成一行、各列串位，于是「按编号切分」这条
    结构线索完全拿不到（真材料上验过：原告证据目录切不准就是这个原因）。
    docx 里表格是结构化的 <w:tbl>，逐行逐格读得出来，所以这条路能做对 —— 每一行
    还原成「编号、名称：内容」这样一句，编号线索就保住了。
    """
    # 只有读 docx 时才需要。缺了就抛一个说得清的错，而不是裸的 ModuleNotFoundError ——
    # 这个 skill 的口径是零第三方依赖，缺一样只该少一种能力，不该让整条路崩掉。
    try:
        from docx import Document
    except ModuleNotFoundError:
        raise DocxUnavailable(
            f"读 .docx 需要 python-docx，当前环境没装：{path}\n"
            f"  装它：pip install python-docx\n"
            f"  或者把这份材料另存为 .txt / .md 再给我 —— 那两种不需要任何库。"
        ) from None
    d = Document(path)
    out = []
    body = d.element.body
    _p, _t = 0, 0
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            if _p < len(d.paragraphs):
                txt = d.paragraphs[_p].text.strip()
                _p += 1
                if txt:
                    # **按段落样式**认标题，不按长度猜。docx 里标题是 Heading 样式，
                    # 这个信号是明确的；按「短且不带句末标点」猜过一版，「证据清单」
                    # 仍然黏在上一句尾部（真材料里出来是「…被告仍未付款。证据清单」），
                    # 因为它长度确实短、但判据落在了猜上。有明确信号时不要猜。
                    _sty = (d.paragraphs[_p - 1].style.name or "")
                    if _sty.startswith("Heading") or _sty.startswith("Title"):
                        txt = txt.rstrip("。；：") + "。"
                    out.append(txt)
        elif tag == "tbl":
            # 表格前若有一段只是表格的标题（如「证据清单」「附表」），它在 docx 里是
            # 独立段落，但被 split_sentences 黏到了上一句尾部（真材料里出来的效果是
            # 「…被告仍未付款。证据清单」）。这类短标题单独成段，与正文分开。
            if _t < len(d.tables):
                tb = d.tables[_t]
                _t += 1
                cells = [[c.text.strip() for c in r.cells] for r in tb.rows]
                head = cells[0] if cells else []
                for r in cells[1:]:
                    # 一行还原成一句，列名带上，避免列与列串位
                    parts = [f"{h}：{v}" for h, v in zip(head, r) if v]
                    if parts:
                        out.append("；".join(parts) + "。")
    return "\n\n".join(out)


def read(path):
    if str(path).lower().endswith(".docx"):
        raw = read_docx(path)
    else:
        raw = open(path, encoding="utf-8").read()
    text, counts = normalise(raw)
    sents = split_sentences(text)
    return {"path": path, "normalised": text, "sentences": sents,
            "normalise_counts": counts}


if __name__ == "__main__":
    r = read(sys.argv[1])
    _raw0 = (read_docx(sys.argv[1]) if sys.argv[1].lower().endswith(".docx")
             else open(sys.argv[1], encoding="utf-8").read())
    print(f"原始 {len(_raw0)} 字 -> "
          f"规范化 {len(r['normalised'])} 字")
    for k, v in r["normalise_counts"].items():
        if v:
            print(f"  {k}: {v} 处")
    s = r["sentences"]
    print(f"切出 {len(s)} 句；长度 中位数 {sorted(len(x) for x in s)[len(s)//2]} "
          f"最长 {max(len(x) for x in s)}")
    if len(sys.argv) > 2:
        json.dump({"sentences": s}, open(sys.argv[2], "w", encoding="utf-8"),
                  ensure_ascii=False, indent=1)
        print(f"句子已写入 {sys.argv[2]}")

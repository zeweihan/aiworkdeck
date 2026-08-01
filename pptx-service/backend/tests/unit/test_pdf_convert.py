"""
[checkba] /api/pdf/* 转换端点回归：

- to-docx：reportlab 生成文本型 PDF → pdf2docx 版式级转换 → python-docx 验证文字与表格保留
- 扫描件拒绝：无文本层 PDF 调 to-docx 必须 400 并指路 ocr-markdown
- ocr-markdown：MinerU 依赖网络/本地服务，这里只验参数校验分支（文件不存在 → 400）
"""
import os

import pytest
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


@pytest.fixture
def text_pdf(tmp_path):
    path = str(tmp_path / "sample.pdf")
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    c.setFont("Helvetica", 14)
    c.drawString(60, h - 80, "Service Agreement")
    c.setFont("Helvetica", 11)
    c.drawString(60, h - 120, "Fee is 20 percent of the total amount.")
    # 画一个简单表格（线框+单元格文本），验证 pdf2docx 的表格识别
    x0, y0, cw, rh = 60, h - 260, 120, 24
    for r in range(3):
        for col in range(2):
            c.rect(x0 + col * cw, y0 - r * rh, cw, rh)
    cells = [("Item", "Price"), ("Alpha", "100"), ("Beta", "200")]
    for r, (a, b) in enumerate(cells):
        c.drawString(x0 + 6, y0 - r * rh + 7, a)
        c.drawString(x0 + cw + 6, y0 - r * rh + 7, b)
    c.showPage()
    c.save()
    return path


@pytest.fixture
def scanned_pdf(tmp_path):
    """无文本层 PDF（空白页即扫描件的最小等价物）"""
    path = str(tmp_path / "scanned.pdf")
    c = canvas.Canvas(path, pagesize=A4)
    c.showPage()
    c.save()
    return path


def test_to_docx_preserves_text_and_table(client, text_pdf, tmp_path):
    out = str(tmp_path / "out.docx")
    resp = client.post('/api/pdf/to-docx', json={'pdf_path': text_pdf, 'output_path': out})
    assert resp.status_code == 200, resp.get_data(as_text=True)
    data = resp.get_json()['data']
    assert data['pages'] == 1
    assert os.path.getsize(out) > 0

    from docx import Document
    doc = Document(out)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(cell.text for t in doc.tables for row in t.rows for cell in row.cells)
    assert "Service Agreement" in all_text
    assert "Fee is 20 percent" in all_text
    assert len(doc.tables) >= 1, "版式级转换应保留表格结构"
    assert "Alpha" in table_text and "200" in table_text


def test_to_docx_rejects_scanned_pdf(client, scanned_pdf, tmp_path):
    resp = client.post('/api/pdf/to-docx',
                       json={'pdf_path': scanned_pdf, 'output_path': str(tmp_path / "o.docx")})
    assert resp.status_code == 400
    assert "ocr-markdown" in resp.get_json()['error']['message']


def test_to_docx_missing_file_rejected(client, tmp_path):
    resp = client.post('/api/pdf/to-docx',
                       json={'pdf_path': str(tmp_path / "nope.pdf"), 'output_path': str(tmp_path / "o.docx")})
    assert resp.status_code == 400


def test_ocr_markdown_missing_file_rejected(client, tmp_path):
    resp = client.post('/api/pdf/ocr-markdown', json={'pdf_path': str(tmp_path / "nope.pdf")})
    assert resp.status_code == 400

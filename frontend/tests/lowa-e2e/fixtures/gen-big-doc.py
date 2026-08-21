#!/usr/bin/env python3
"""大文档基线组的夹具生成器（dev-board#108，配方见
docs/superpowers/specs/2026-08-21-dd-scale-stability-inventory.md 附录 A）。

产出：150 页 docx。每页 1 个二级标题 + 4 段约 220 字中文正文 + 分页符；随机 30 页
各插一张 12x5 `Table Grid` 表；随机 20 页各插一张 900x600 噪声 JPEG（quality 75）。
随机种子固定（20260821），所以每次生成的结构完全一致——基线数字才有可比性。

每页首段各含一次「目标公司」（全文恰好 150 处），给 find_replace 150 命中的计时项用。

输出：$TMPDIR/awd-big-doc/big.docx（不入库；可用 --out 改路径）。
依赖：python3 -m pip install --user python-docx pillow
"""
import argparse
import io
import os
import random
import sys
import tempfile

try:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Inches
    from PIL import Image
except ImportError as e:  # pragma: no cover - 环境提示
    sys.stderr.write(
        '缺少依赖 (%s)。请先执行：\n  python3 -m pip install --user python-docx pillow\n' % e)
    sys.exit(2)

PAGES = 150
PARAS_PER_PAGE = 4
TABLE_PAGES = 30
IMAGE_PAGES = 20
SEED = 20260821
HIT_TOKEN = '目标公司'

WORDS = (
    '根据', '公司章程', '股东会', '决议', '本所律师', '核查', '截至', '报告期', '未发现',
    '重大', '违法违规', '情形', '相关', '资产', '权属', '清晰', '不存在', '抵押', '质押',
    '查封', '冻结', '对外担保', '诉讼', '仲裁', '行政处罚', '劳动合同', '社会保险', '住房公积金',
    '缴纳', '税务', '合规', '知识产权', '商标', '专利', '软件著作权', '取得', '登记证书',
    '有效期内', '土地使用权', '房屋', '租赁', '协议', '履行', '经营范围', '许可', '资质',
    '关联交易', '定价', '公允', '同业竞争', '承诺', '董事', '监事', '高级管理人员', '任职',
    '资格', '符合', '法律', '法规', '规范性文件', '规定', '经营', '主要', '客户', '供应商',
    '采购', '销售', '合同', '条款', '风险', '提示', '建议', '整改', '完成', '补充', '说明',
)


def make_paragraph_text(rnd, with_token):
    out = []
    length = 0
    while length < 220:
        w = rnd.choice(WORDS)
        out.append(w)
        length += len(w)
        if rnd.random() < 0.12:
            out.append('，')
        if rnd.random() < 0.05:
            out.append('。')
    text = ''.join(out).rstrip('，。') + '。'
    if with_token:
        # 固定插在段首之后，保证每段恰好一处命中（WORDS 里没有这个词）
        text = '经核查，' + HIT_TOKEN + text
    return text


def noise_jpeg(rnd):
    # 随机字节噪声图：JPEG 压不动，单张约 330KB，刻意把文件撑大
    raw = bytes(rnd.getrandbits(8) for _ in range(900 * 600 * 3))
    img = Image.frombytes('RGB', (900, 600), raw)
    buf = io.BytesIO()
    img.save(buf, format='JPEG', quality=75)
    buf.seek(0)
    return buf


def build(out_path):
    rnd = random.Random(SEED)
    table_pages = set(rnd.sample(range(PAGES), TABLE_PAGES))
    image_pages = set(rnd.sample(range(PAGES), IMAGE_PAGES))
    doc = Document()
    doc.add_heading('尽职调查报告（大文档基线夹具）', level=1)
    for page in range(PAGES):
        doc.add_heading('第%d节 核查事项' % (page + 1), level=2)
        for i in range(PARAS_PER_PAGE):
            doc.add_paragraph(make_paragraph_text(rnd, with_token=(i == 0)))
        if page in table_pages:
            table = doc.add_table(rows=12, cols=5)
            table.style = 'Table Grid'
            for r in range(12):
                for c in range(5):
                    table.cell(r, c).text = ('项目' if r == 0 else '%d' % rnd.randint(100, 99999)) if c else ('第%d行' % r)
        if page in image_pages:
            doc.add_picture(noise_jpeg(rnd), width=Inches(5.5))
        if page < PAGES - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser()
    default_out = os.path.join(tempfile.gettempdir(), 'awd-big-doc', 'big.docx')
    ap.add_argument('--out', default=default_out)
    ap.add_argument('--force', action='store_true', help='已存在也重新生成')
    args = ap.parse_args()
    if os.path.exists(args.out) and not args.force:
        print(args.out)
        return
    build(args.out)
    print(args.out)


if __name__ == '__main__':
    main()
